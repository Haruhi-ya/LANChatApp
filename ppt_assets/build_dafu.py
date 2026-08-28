# -*- coding: utf-8 -*-
"""生成答辩记录文档（docx）"""
import sys
sys.stdout.reconfigure(encoding='utf-8')

import docx
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import os

doc = docx.Document()

# 全局字体
style = doc.styles['Normal']
style.font.name = '宋体'
style.font.size = Pt(12)
style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

def set_ea(run, name='宋体'):
    run.font.name = name
    run._element.rPr.rFonts.set(qn('w:eastAsia'), name)

def para(text, size=12, bold=False, align=None, color=None, space_after=6):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.size = Pt(size)
    r.bold = bold
    if color:
        r.font.color.rgb = color
    if align:
        p.alignment = align
    set_ea(r)
    p.paragraph_format.space_after = Pt(space_after)
    return p

# ============ 封面信息 ============
para('答 辩 记 录', size=22, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=18)
para('项目名称：局域网聊天室（即时通讯工具）', size=13, space_after=4)
para('小组成员：赵翔（25112002136）、胡鸣（25112002139）、茹桂堂（25112002124）', size=13, space_after=4)
para('汇报方式：一名同学主汇报，其他两名同学补充', size=13, space_after=4)
para('答辩时间：2026 年 ____ 月 ____ 日', size=13, space_after=4)
para('答辩地点：______________', size=13, space_after=12)

para('一、答辩过程简述', size=15, bold=True, space_after=8)
para('本次答辩由赵翔担任主汇报人，围绕项目背景、软件需求、概要设计、详细设计、系统实现与系统测试六个方面进行汇报，'
     '汇报时长控制在 5 分钟以内。汇报结束后，胡鸣补充了数据库设计与数据访问层的内容，茹桂堂补充了界面实现与交互设计的内容。'
     '随后评审老师就系统设计、实现细节与测试方法等提出若干问题，小组成员现场作答，记录整理如下。', size=12.5, space_after=12)

para('二、老师提问与回答记录', size=15, bold=True, space_after=8)

# 空记录表
tbl = doc.add_table(rows=6, cols=4)
tbl.style = 'Table Grid'
tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
widths = [Cm(1.5), Cm(5.5), Cm(6.5), Cm(2.0)]
headers = ['题号', '老师提问内容', '小组回答要点', '作答人']
for j, h in enumerate(headers):
    c = tbl.rows[0].cells[j]
    c.text = ''
    p = c.paragraphs[0]
    r = p.add_run(h)
    r.bold = True
    r.font.size = Pt(11)
    set_ea(r)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
for i in range(1, 6):
    for j in range(4):
        c = tbl.rows[i].cells[j]
        c.text = ''
        if j == 0:
            p = c.paragraphs[0]
            r = p.add_run(str(i))
            r.font.size = Pt(11)
            set_ea(r)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
for row in tbl.rows:
    for j, w in enumerate(widths):
        row.cells[j].width = w

para('', size=6, space_after=2)
para('三、答辩预判问题与参考回答（供现场准备参考）', size=15, bold=True, space_after=8)

qa = [
    ('Q1：为什么采用自定义文本行协议，而不是 JSON 或对象序列化？',
     '文本行协议轻量、可读、便于调试，与 Java 对象流相比没有版本兼容问题，与 JSON 相比解析开销更小。'
     '关键设计是"内容字段固定放最后"，解析时用 indexOf(":") 只切前面固定数量的字段，消息内容可以安全包含冒号；'
     '用户名禁止冒号与逗号，时间戳使用 epoch 毫秒纯数字，规避了格式自带冒号破坏字段的问题。'),
    ('Q2：消息撤回是如何实现的？如何防止伪造撤回？',
     '每条消息入库时由服务端生成全局唯一的 UUID 作为消息 ID（client_id 列），客户端构造的任何 ID 都会被 isValidMsgId 格式校验拦截。'
     '撤回时在同一 synchronized 方法内完成权限校验（普通用户只能撤回自己的消息，管理员可撤回任意公共消息）与两分钟时间窗口校验，'
     '时间以服务端入库时间为准，不信任客户端时间；通过后向全体在线客户端广播撤回事件，各端把消息替换为撤回提示。'),
    ('Q3：私聊记录为什么要设计成"双份行模型"？清空记录为什么只影响一方？',
     '私聊消息表每条消息同时写入归属双方各一行（owner 视角字段），双方各自持有独立副本。'
     '这样任意一方清空记录时只删除自己名下的行，对方保存的副本不受影响，实现了"记录清空只影响自己"的需求，'
     '同时未读状态（is_read）天然按接收方视角独立存储，未读统计也直接基于该模型查询。'),
    ('Q4：头像为什么采用"按需拉取"而不是随用户列表广播？',
     '头像数据较大（最大 128KB，Base64 后更大），若随用户列表广播会造成大量冗余流量。'
     '按需拉取配合内存缓存与负缓存避免重复请求；头像变更时服务端广播通知，各端清缓存后重新拉取。'
     '缓存采用版本号机制：拉取发起时记录版本号，响应到达时若版本已变化则丢弃陈旧数据，保证用户快速连续更换头像时各端最终显示最新头像。'),
    ('Q5：换行注入攻击是什么？为什么客户端过滤不够，必须在服务端防御？',
     '攻击者在消息内容中混入 \\r\\n，服务端 readLine() 会按行读取，把一条消息拆成两条命令，例如"粘贴文本\\r\\nCLEARPUBLIC"变成一条消息加一条清空命令，可能造成越权操作。'
     '客户端 sendLine() 虽然会过滤换行，但任何客户端都可以被绕过（用原始 socket 直接发送），因此服务端 sanitize() 在清洗阶段再次过滤、校验与限长，'
     '测试中我们用原始 socket 构造含 \\r\\n 的报文验证了服务端防线确实生效——服务端校验才是真正的安全边界。'),
    ('Q6：如何保证历史回放与实时消息的顺序一致？',
     '服务端消息管道为"先入库、后广播"：入库失败则不广播，保证数据库中的历史记录与实时广播顺序完全一致。'
     '客户端历史加载采用 BEGIN→ITEM×N→END 三段式，历史加载期间到达的实时消息先缓冲，等 HISTEND 到达后再按序补渲染，'
     '避免并发写入导致的重复或乱序；任何异常分支也会发送 END，客户端不会永远停留在加载状态。'),
    ('Q7：服务端是如何处理并发与断连的？',
     '服务端采用一连接一线程模型，主线程只做 accept；在线客户端表使用 ConcurrentHashMap，其弱一致性迭代器保证广播遍历过程中并发删除安全。'
     '断开清理统一收敛到 disconnect()：条件删除（remove(key, value)）保证只删除自己的连接，nickname 判空保证幂等，'
     '正常断开与异常断开同时触发时不会重复广播。客户端网络层使用独立接收线程（daemon），回调经 SwingUtilities.invokeLater 切回 EDT 更新界面，网络线程不触碰 Swing 组件。'),
    ('Q8：未读数是如何统计的？客户端为什么不自己计数？',
     '未读数由服务端基于私聊表双份行模型查询 is_read=0 的记录数得到，是服务端权威值。'
     '客户端收到私聊消息时不自行累加，而是使用服务端随消息下发的未读数，避免与登录时的 UNREAD 未读汇总互相覆盖不一致；'
     '撤回未读消息后服务端会向接收方补发新的未读汇总，纠正红点数。'),
    ('Q9：@提醒功能如何实现？中文用户名会不会被错误匹配？',
     '消息清洗入库广播后，服务端扫描内容中是否包含"@在线用户名"。@ 匹配使用 Unicode 类别边界（\\p{L}\\p{N}）而不是 Java 的 \\w，'
     '因为 \\w 在 Java 中默认不支持中文，直接用会导致中文用户名被错误切分；同一消息对同一人只提醒一次。'
     '客户端收到消息广播后渲染气泡，@提醒事件（ATMSG）在消息广播之后发送，保证先渲染再闪窗提示。'),
    ('Q10：测试为什么采用协议级黑盒测试，而不是通过 GUI 操作测试？',
     '协议级黑盒测试用 Python 原始 socket 直接向服务端发送协议命令并断言响应与数据库状态，可以精确控制输入、批量并发、'
     '构造攻击报文（如换行注入、伪造魔数、伪造消息 ID），这些是 GUI 测试难以稳定复现的；'
     '同时协议是服务端与客户端唯一的契约，协议级测试能直接验证服务端独立防线。'
     '测试覆盖六大功能域 34 个用例、37 项断言全部通过，含 5 个攻击场景，测试后核查数据库零残留。'),
]
for q, a in qa:
    p = para('【问题】' + q, size=12, bold=True, space_after=2)
    para('【回答】' + a, size=12, space_after=10)

para('', size=6)
para('记录人：______________        审核人：______________', size=12, space_after=4)

OUT = os.path.join(os.path.dirname(os.path.abspath(__file__)), '..', '答辩记录-局域网聊天室.docx')
doc.save(OUT)
print('已生成:', os.path.abspath(OUT))
