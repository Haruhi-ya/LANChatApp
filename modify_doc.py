# -*- coding: utf-8 -*-
"""
修改「详细设计及测试报告」文档：
1. 所有表格内容水平/垂直居中
2. 所有表格改为三线式（上下 1.5pt 粗线 + 表头行下 0.75pt 细线，去掉其余边框）
3. 标题样式颜色改为黑色
4. 封面后插入目录（TOC 域）
"""
import sys, shutil, os
sys.stdout.reconfigure(encoding='utf-8')

from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import RGBColor, Pt

SRC = '详细设计及测试报告_赵翔_25112002136.docx'

# ---------- 0. 备份 ----------
bak = SRC.replace('.docx', '_bak.docx')
shutil.copy2(SRC, bak)
print(f'已备份 -> {bak}')

doc = Document(SRC)

# ---------- 1. 标题颜色改为黑色 ----------
for name in ('Heading 1', 'Heading 2', 'Heading 3'):
    try:
        doc.styles[name].font.color.rgb = RGBColor(0, 0, 0)
        print(f'样式 {name} 颜色已改为黑色')
    except Exception as e:
        print(f'样式 {name} 修改失败: {e}')

# ---------- 2. 表格处理 ----------
def set_tbl_borders(table):
    """表格级边框：top/bottom 1.5pt，其余 none（三线式）"""
    tblPr = table._tbl.tblPr
    for old in tblPr.findall(qn('w:tblBorders')):
        tblPr.remove(old)
    borders = OxmlElement('w:tblBorders')
    for edge, val, sz in (('top', 'single', 12), ('bottom', 'single', 12)):
        el = OxmlElement(f'w:{edge}')
        el.set(qn('w:val'), val)
        el.set(qn('w:sz'), str(sz))
        el.set(qn('w:space'), '0')
        el.set(qn('w:color'), '000000')
        borders.append(el)
    for edge in ('left', 'right', 'insideH', 'insideV'):
        el = OxmlElement(f'w:{edge}')
        el.set(qn('w:val'), 'none')
        borders.append(el)
    tblLayout = tblPr.find(qn('w:tblLayout'))
    if tblLayout is not None:
        tblLayout.addprevious(borders)
    else:
        tblPr.append(borders)

def set_header_row_bottom(table):
    """表头行（第一行）下边框 0.75pt"""
    for tc in table.rows[0]._tr.findall(qn('w:tc')):
        tcPr = tc.get_or_add_tcPr()
        for old in tcPr.findall(qn('w:tcBorders')):
            tcPr.remove(old)
        tcb = OxmlElement('w:tcBorders')
        bottom = OxmlElement('w:bottom')
        bottom.set(qn('w:val'), 'single')
        bottom.set(qn('w:sz'), '6')      # 0.75pt
        bottom.set(qn('w:space'), '0')
        bottom.set(qn('w:color'), '000000')
        tcb.append(bottom)
        vAlign = tcPr.find(qn('w:vAlign'))
        if vAlign is not None:
            vAlign.addprevious(tcb)
        else:
            tcPr.append(tcb)

def center_all_cells(table):
    """所有单元格：水平居中 + 垂直居中"""
    for tc in table._tbl.iter(qn('w:tc')):
        tcPr = tc.get_or_add_tcPr()
        vAlign = tcPr.find(qn('w:vAlign'))
        if vAlign is None:
            vAlign = OxmlElement('w:vAlign')
            tcPr.append(vAlign)
        vAlign.set(qn('w:val'), 'center')
        for p in tc.findall(qn('w:p')):
            pPr = p.find(qn('w:pPr'))
            if pPr is None:
                pPr = OxmlElement('w:pPr')
                p.insert(0, pPr)
            jc = pPr.find(qn('w:jc'))
            if jc is None:
                jc = OxmlElement('w:jc')
                pPr.append(jc)
            jc.set(qn('w:val'), 'center')

for i, table in enumerate(doc.tables):
    set_tbl_borders(table)
    set_header_row_bottom(table)
    center_all_cells(table)
    print(f'表格[{i}] 三线式+居中完成 ({len(table.rows)}行x{len(table.columns)}列)')

# ---------- 3. 插入目录 ----------
def make_para():
    return OxmlElement('w:p')

def add_run(p, text=None, bold=False, size=None):
    r = OxmlElement('w:r')
    if bold or size:
        rPr = OxmlElement('w:rPr')
        if bold:
            rPr.append(OxmlElement('w:b'))
        if size:
            sz = OxmlElement('w:sz')
            sz.set(qn('w:val'), str(int(size * 2)))  # half-points
            rPr.append(sz)
        r.append(rPr)
    if text is not None:
        t = OxmlElement('w:t')
        t.text = text
        t.set(qn('xml:space'), 'preserve')
        r.append(t)
    p.append(r)
    return p

body = doc.element.body
# 封面表格后第一个段落（空段 body[12]）作为插入锚点之后的位置
anchor = doc.tables[0]._tbl.getnext()

# 3.1 目录标题：居中、加粗、16pt（与 Heading 1 同字号）
title_p = make_para()
pPr = OxmlElement('w:pPr')
jc = OxmlElement('w:jc')
jc.set(qn('w:val'), 'center')
pPr.append(jc)
title_p.append(pPr)
add_run(title_p, '目　　录', bold=True, size=16)

# 3.2 TOC 域段落
toc_p = make_para()
# begin
r = OxmlElement('w:r')
fld = OxmlElement('w:fldChar')
fld.set(qn('w:fldCharType'), 'begin')
fld.set(qn('w:dirty'), 'true')
r.append(fld)
toc_p.append(r)
# instrText
r = OxmlElement('w:r')
instr = OxmlElement('w:instrText')
instr.text = ' TOC \\o "1-3" \\h \\z \\u '
instr.set(qn('xml:space'), 'preserve')
r.append(instr)
toc_p.append(r)
# separate
r = OxmlElement('w:r')
fld = OxmlElement('w:fldChar')
fld.set(qn('w:fldCharType'), 'separate')
r.append(fld)
toc_p.append(r)
# 占位文本
r = OxmlElement('w:r')
t = OxmlElement('w:t')
t.text = '（请在 Word 中按 Ctrl+A 后按 F9，或右键此处选择"更新域"生成目录）'
r.append(t)
toc_p.append(r)
# end
r = OxmlElement('w:r')
fld = OxmlElement('w:fldChar')
fld.set(qn('w:fldCharType'), 'end')
r.append(fld)
toc_p.append(r)

# 3.3 分页段落（目录后正文另起一页）
page_p = make_para()
r = OxmlElement('w:r')
br = OxmlElement('w:br')
br.set(qn('w:type'), 'page')
r.append(br)
page_p.append(r)

anchor.addprevious(title_p)
anchor.addprevious(toc_p)
anchor.addprevious(page_p)
print('目录已插入（封面表之后）')

# ---------- 4. 保存 ----------
try:
    doc.save(SRC)
    print(f'已保存 -> {SRC}')
except PermissionError:
    out = SRC.replace('.docx', '_new.docx')
    doc.save(out)
    print(f'原文件被占用（可能 Word 正打开），已保存 -> {out}')
