# -*- coding: utf-8 -*-
"""验证修改后的文档"""
import sys, re
sys.stdout.reconfigure(encoding='utf-8')

import docx
from docx.oxml.ns import qn

doc = docx.Document('详细设计及测试报告_赵翔_25112002136_new.docx')

# 1. 标题颜色
print("=== 标题样式颜色 ===")
for name in ('Heading 1', 'Heading 2', 'Heading 3'):
    s = doc.styles[name]
    color = s.element.find(f'{qn("w:rPr")}/{qn("w:color")}')
    print(f"{name}: color={color.get(qn('w:val')) if color is not None else None}")

# 2. 表格边框与居中
print("\n=== 表格检查 ===")
ok = True
for i, t in enumerate(doc.tables):
    tblPr = t._tbl.tblPr
    tb = tblPr.find(qn('w:tblBorders'))
    top = tb.find(qn('w:top')) if tb is not None else None
    bottom = tb.find(qn('w:bottom')) if tb is not None else None
    insideV = tb.find(qn('w:insideV')) if tb is not None else None
    top_ok = top is not None and top.get(qn('w:sz')) == '12'
    bottom_ok = bottom is not None and bottom.get(qn('w:sz')) == '12'
    inside_ok = insideV is not None and insideV.get(qn('w:val')) == 'none'
    # 表头行底边框
    hdr_tc = t.rows[0]._tr.findall(qn('w:tc'))[0]
    tcb = hdr_tc.find(f'{qn("w:tcPr")}/{qn("w:tcBorders")}')
    hdr_ok = tcb is not None and tcb.find(qn('w:bottom')).get(qn('w:sz')) == '6'
    # 单元格居中抽查：最后一个单元格
    last_tc = list(t._tbl.iter(qn('w:tc')))[-1]
    valign = last_tc.find(f'{qn("w:tcPr")}/{qn("w:vAlign")}')
    v_ok = valign is not None and valign.get(qn('w:val')) == 'center'
    jc = last_tc.find(f'{qn("w:p")}/{qn("w:pPr")}/{qn("w:jc")}')
    h_ok = jc is not None and jc.get(qn('w:val')) == 'center'
    status = 'OK' if all([top_ok, bottom_ok, inside_ok, hdr_ok, v_ok, h_ok]) else 'FAIL'
    if status == 'FAIL':
        ok = False
        print(f"表{i}: top={top_ok} bottom={bottom_ok} insideV={inside_ok} hdrBottom={hdr_ok} vAlign={v_ok} jc={h_ok}")
print("所有表格格式:", "全部 OK" if ok else "存在问题")

# 3. 目录域
xml = doc.element.body.xml
has_toc = 'TOC \\o' in xml and 'fldChar' in xml
print("\n=== TOC 域:", "已插入" if has_toc else "缺失")

# 4. 文档结构：封面表后新插入的段落
body = doc.element.body
print("\n=== 封面表后的段落顺序 ===")
tbl0 = doc.tables[0]._tbl
el = tbl0.getnext()
for _ in range(6):
    if el is None:
        break
    if el.tag.endswith('}p'):
        texts = el.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}t')
        print("p:", ''.join(t.text or '' for t in texts)[:60] or '(空/分页)')
    else:
        print(el.tag.split('}')[1])
    el = el.getnext()
